# export_utils.py

import io
from docx import Document
from docx.shared import RGBColor
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.enums import TA_LEFT
from reportlab.lib import colors


# ---------------- DOCX Export ---------------- #
def export_quiz_docx(quiz_data, with_answers=True):
    """
    Export quiz to DOCX.
    Styling:
      - Question: blue, bold
      - Options: normal
      - Answer: green, bold
      - Explanation: bold
    """
    doc = Document()

    for idx, q in enumerate(quiz_data, 1):
        # Question
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {q['question']}")
        run.bold = True
        run.font.color.rgb = RGBColor(53, 94, 143)  # Blue

        # Options
        for key, val in q["options"].items():
            doc.add_paragraph(f"{key}) {val}")

        if with_answers:
            # Answer
            ans_par = doc.add_paragraph()
            ans_run = ans_par.add_run(f"Answer: {q['answer']}) {q['options'].get(q['answer'], '')}")
            ans_run.bold = True
            ans_run.font.color.rgb = RGBColor(0, 158, 71)  # Green

            # Explanation
            exp_par = doc.add_paragraph()
            exp_run = exp_par.add_run(f"Explanation: {q.get('explanation','')}")
            exp_run.bold = True

        doc.add_paragraph("")  # blank line

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ---------------- PDF Export ---------------- #
def export_quiz_pdf_bytes(quiz_data, with_answers=True):
    """
    Export quiz to PDF with styles.
    Styling:
      - Question: blue (#355E8F), bold
      - Options: normal
      - Answer: green (#009E47), bold
      - Explanation: bold (black)
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=36, leftMargin=36,
        topMargin=36, bottomMargin=36
    )

    styles = getSampleStyleSheet()

    question_style = ParagraphStyle(
        name="QuestionStyle",
        parent=styles["Heading4"],
        fontName="Helvetica-Bold",
        fontSize=12,
        leading=14,
        textColor=colors.HexColor("#355E8F"),
        spaceAfter=4,
        alignment=TA_LEFT
    )

    option_style = ParagraphStyle(
        name="OptionStyle",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=12,
        leftIndent=8,
        spaceAfter=2
    )

    answer_style = ParagraphStyle(
        name="AnswerStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=10,
        leading=12,
        textColor=colors.HexColor("#009E47"),
        spaceBefore=6,
        spaceAfter=2
    )

    explanation_style = ParagraphStyle(
        name="ExplanationStyle",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=9,
        leading=11,
        textColor=colors.black,
        spaceAfter=8
    )

    story = []
    for idx, q in enumerate(quiz_data, 1):
        # Question
        story.append(Paragraph(f"{idx}. {q['question']}", question_style))

        # Options
        for opt_key, opt_text in q["options"].items():
            story.append(Paragraph(f"{opt_key}) {opt_text}", option_style))

        if with_answers:
            ans_text = f"{q['answer']}) {q['options'].get(q['answer'], '')}"
            story.append(Paragraph(f"Answer: {ans_text}", answer_style))
            story.append(Paragraph(f"Explanation: {q.get('explanation','')}", explanation_style))

        story.append(Spacer(1, 10))

    doc.build(story)
    buffer.seek(0)
    return buffer



# from docx import Document
# from docx.shared import RGBColor
# import io
# from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
# from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
# from reportlab.lib.enums import TA_LEFT
# from reportlab.lib import colors


# # DOCX Export
# def export_quiz_docx(quiz_data, with_answers=True):
#     doc = Document()

#     for idx, q in enumerate(quiz_data, 1):
#         # Question styled (blue + bold)
#         p = doc.add_paragraph()
#         run = p.add_run(f"{idx}. {q['question']}")
#         run.bold = True
#         run.font.color.rgb = RGBColor(53, 94, 143)  # blue

#         # Options
#         for key, val in q["options"].items():
#             doc.add_paragraph(f"{key}) {val}")

#         if with_answers:
#             # Answer styled (green + bold)
#             p_ans = doc.add_paragraph()
#             run_ans = p_ans.add_run(f"Answer: {q['answer']}")
#             run_ans.bold = True
#             run_ans.font.color.rgb = RGBColor(0, 158, 71)  # green

#             # Explanation styled (bold only)
#             p_exp = doc.add_paragraph()
#             run_exp = p_exp.add_run(f"Explanation: {q['explanation']}")
#             run_exp.bold = True
#             run_exp.italic = True

#         doc.add_paragraph("")  # blank line

#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer


# # PDF Export
# def export_quiz_pdf(quiz_data, with_answers=True, filename="quiz.pdf"):
#     styles = getSampleStyleSheet()

#     # Custom styles
#     question_style = ParagraphStyle(
#         name="QuestionStyle",
#         parent=styles["Heading4"],
#         fontName="Helvetica-Bold",
#         fontSize=12,
#         textColor=colors.Color(53/255, 94/255, 143/255),  # blue
#         alignment=TA_LEFT
#     )

#     option_style = styles["Normal"]

#     answer_style = ParagraphStyle(
#         name="AnswerStyle",
#         parent=styles["Normal"],
#         fontName="Helvetica-Bold",
#         fontSize=11,
#         textColor=colors.Color(0/255, 158/255, 71/255),  # green
#     )

#     explanation_style = ParagraphStyle(
#         name="ExplanationStyle",
#         parent=styles["Normal"],
#         fontName="Helvetica-Bold",
#         fontSize=10,
#     )

#     story = []
#     for idx, q in enumerate(quiz_data, 1):
#         # Question
#         story.append(Paragraph(f"{idx}. {q['question']}", question_style))

#         # Options
#         for opt_key, opt_text in q["options"].items():
#             story.append(Paragraph(f"{opt_key}) {opt_text}", option_style))

#         if with_answers:
#             # Answer
#             story.append(Paragraph(f"Answer: {q['answer']}", answer_style))
#             # Explanation
#             story.append(Paragraph(f"Explanation: {q['explanation']}", explanation_style))

#         story.append(Spacer(1, 12))  # space

#     doc = SimpleDocTemplate(filename)
#     doc.build(story)
#     return filename